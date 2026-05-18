#!/usr/bin/env python3
"""Build/render weekly Bee actionable reports.

This script does not call AI models. It prepares source digest material and
renders an agent-written actionable-report.md to DOCX/PDF. The OpenClaw cron
agentTurn does the reasoning/synthesis from the source files.
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

ROOT = Path(__file__).resolve().parents[1]


def read_json(path: Path, default):
    try:
        return json.loads(path.read_text())
    except Exception:
        return default


def first_text(obj: dict, keys: Iterable[str]) -> str:
    for key in keys:
        val = obj.get(key)
        if isinstance(val, str) and val.strip():
            return val.strip()
    return ""


def strip_md(text: str) -> str:
    text = re.sub(r"[`*_#]", "", text)
    return text.strip()


def clean_line(text: Any, limit: int = 220) -> str:
    line = " ".join(str(text or "").split())
    if len(line) <= limit:
        return line
    return line[: limit - 1].rstrip() + "..."


def local_date(ms: Any) -> str:
    if not isinstance(ms, (int, float)) or ms <= 0:
        return ""
    return datetime.fromtimestamp(ms / 1000).strftime("%Y-%m-%d")


def item_text(item: dict) -> str:
    return first_text(item, ["text", "content", "title", "summary", "short_summary"])


def is_open_todo(item: dict) -> bool:
    status = str(item.get("status") or item.get("state") or "").lower()
    if item.get("completed") is True:
        return False
    return status not in {"done", "complete", "completed", "archived", "deleted"}


def classify(text: str) -> str:
    lowered = text.lower()
    buckets = [
        ("Family / Home", ["amanda", "franklin", "may", "kid", "school", "vbs", "camp", "house", "home", "clean", "grill", "water bottle", "marriage"]),
        ("Work / UAB", ["uab", "heal", "gimpop", "newsletter", "grant", "retreat", "faculty", "department", "work"]),
        ("Church / Sermon", ["church", "sermon", "brandon", "av", "venue", "podcast", "transistor", "audio"]),
        ("Health / Training", ["run", "training", "doctor", "appointment", "med", "sleep", "rash", "health"]),
    ]
    for label, needles in buckets:
        if any(needle in lowered for needle in needles):
            return label
    return "General"


def validate_action_report(path: Path) -> None:
    text = path.read_text(errors="replace")
    forbidden_headings = [
        "## Daily summaries",
        "## Conversation signals",
        "## Transcript files available",
        "## Facts created/updated in window",
        "## Todos created/alarmed in window",
    ]
    for heading in forbidden_headings:
        if heading.lower() in text.lower():
            raise SystemExit(f"Action report includes source-dump heading: {heading}")


def collect_action_candidates(source_dir: Path) -> dict[str, list[str]]:
    buckets: dict[str, list[str]] = {
        "Family / Home": [],
        "Work / UAB": [],
        "Church / Sermon": [],
        "Health / Training": [],
        "General": [],
    }

    def add(raw: str, prefix: str = "") -> None:
        line = clean_line(raw)
        if not line:
            return
        rendered = f"{prefix}{line}" if prefix else line
        bucket = classify(line)
        if rendered not in buckets[bucket]:
            buckets[bucket].append(rendered)

    todos = read_json(source_dir / "todos-week.json", [])
    for item in todos:
        if is_open_todo(item):
            due = local_date(item.get("alarm_at") or item.get("due_at") or item.get("created_at"))
            add(item_text(item), f"{due}: " if due else "")

    action_words = re.compile(r"\b(need|needs|todo|to do|follow up|call|email|ask|send|schedule|register|confirm|prepare|fix|clean|buy|replace|draft|publish|upload|decide|approval)\b", re.I)
    for filename, max_items in [
        ("facts-week.json", 200),
        ("journals-week.json", 80),
        ("daily-week.json", 14),
        ("conversations-week.json", 100),
    ]:
        for item in read_json(source_dir / filename, [])[:max_items]:
            text = item_text(item)
            if action_words.search(text):
                add(text)

    return buckets


def build(args: argparse.Namespace) -> None:
    report_root = Path(args.report_root).resolve()
    source_dir = Path(args.source_dir).resolve()
    report_root.mkdir(parents=True, exist_ok=True)

    window = read_json(report_root / "window.json", {})
    counts = read_json(source_dir / "export-counts.json", {})
    buckets = collect_action_candidates(source_dir)
    all_actions = [item for values in buckets.values() for item in values]
    since = str(window.get("since_local", "?")).split("T")[0]
    until = str(window.get("until_local", "?")).split("T")[0]
    tz = window.get("timezone", "America/Chicago")

    lines: list[str] = [
        f"# Weekly Bee Action Report — {window.get('report_date', report_root.name)}",
        "",
        f"Window: {since} → {until} {tz}",
        f"Source counts: {json.dumps(counts or {'conversations': 0, 'daily': 0, 'facts': 0, 'todos': 0, 'journals': 0}, sort_keys=True)}",
        "",
        "## Executive Summary",
        "",
    ]
    if all_actions:
        lines.append(f"Found {len(all_actions)} local action candidates across Bee todos, facts, journals, daily summaries, and conversation metadata. This report is intentionally actionable-only; transcripts and raw summaries stay in the local source folder for audit.")
    else:
        lines.append("No Bee action candidates were found in this window. Nothing should be inferred from an empty export; check Bee export warnings if counts look unexpectedly low.")

    lines.extend(["", "## Top Pain Points", ""])
    if all_actions:
        for label, items in buckets.items():
            if items:
                lines.append(f"- {label}: {clean_line(items[0], 180)}")
    else:
        lines.append("- No pain points detected from this window.")

    lines.extend(["", "## Top 3 Actions This Week", ""])
    for idx, item in enumerate(all_actions[:3], 1):
        lines.append(f"{idx}. {item}")
    if not all_actions:
        lines.append("1. No action needed from Bee data for this window.")

    sections = [
        ("Family / Home / Relationship Action Plan", "Family / Home"),
        ("Work / UAB Action Plan", "Work / UAB"),
        ("Church / Sermon Action Plan", "Church / Sermon"),
        ("Health / Training Action Plan", "Health / Training"),
    ]
    for heading, bucket in sections:
        lines.extend(["", f"## {heading}", ""])
        items = buckets.get(bucket, [])[:8]
        if items:
            lines.extend(f"- {item}" for item in items)
        else:
            lines.append("- No specific local action surfaced.")

    lines.extend(["", "## Action Packets, Scripts, and Checklists", ""])
    if all_actions:
        lines.extend([
            "- Convert repeated family logistics into short checklists before the next busy morning.",
            "- Keep project, church, and publishing workflows in local runbooks so the next pass is procedural.",
            "- Draft outbound messages locally first; send only after explicit approval.",
        ])
    else:
        lines.append("- No packets needed for an empty window.")

    decision_words = re.compile(r"\b(decide|approval|approve|register|pay|payment|send|contact|call|email|publish|schedule)\b", re.I)
    decisions = [item for item in all_actions if decision_words.search(item)]
    lines.extend(["", "## Chris Must Decide or Do", ""])
    if decisions:
        lines.extend(f"- {item}" for item in decisions[:8])
    else:
        lines.append("- No Chris-only decision surfaced.")

    lines.extend(["", "## Assistant Can Do Next Without External Side Effects", ""])
    lines.extend([
        "- Draft checklists, scripts, and message templates locally.",
        "- Update local planning files from verified source material.",
        "- Re-render this report after edits without re-exporting Bee data.",
    ])

    lines.extend(["", "## Approval Needed", ""])
    lines.extend([
        "- Sending texts, emails, or public posts.",
        "- Registering a child, accepting guardian/legal terms, or making payments.",
        "- Calendar changes, provider contacts, school/church contacts, or publishing audio.",
    ])

    lines.extend(["", "## Blockers / Caveats", ""])
    lines.extend([
        "- Bee speaker labels and summaries are useful signals, not ground truth for sensitive topics.",
        "- This report intentionally excludes conversation-summary appendices and raw transcript dumps.",
    ])

    out = report_root / "actionable-report.md"
    out.write_text("\n".join(lines).rstrip() + "\n")
    validate_action_report(out)
    print(out)


def digest(args: argparse.Namespace) -> None:
    report_root = Path(args.report_root).resolve()
    source_dir = Path(args.source_dir).resolve()
    report_root.mkdir(parents=True, exist_ok=True)

    window = read_json(report_root / "window.json", {})
    counts = read_json(source_dir / "export-counts.json", {})
    conversations = read_json(source_dir / "conversations-week.json", [])
    facts = read_json(source_dir / "facts-week.json", [])
    todos = read_json(source_dir / "todos-week.json", [])
    journals = read_json(source_dir / "journals-week.json", [])
    daily = read_json(source_dir / "daily-week.json", [])

    lines: list[str] = []
    lines.append(f"# Bee Weekly Source Digest — {window.get('report_date', report_root.name)}")
    lines.append("")
    lines.append(f"Window: {window.get('since_local', '?')} → {window.get('until_local', '?')} {window.get('timezone', 'America/Chicago')}")
    lines.append(f"Counts: {counts}")
    lines.append("")
    lines.append("## Instructions for report synthesis")
    lines.append("Generate ONLY actionable report pages. Do not include a conversation-summary appendix. Do not quote long transcripts. Use transcripts/source files only as evidence.")
    lines.append("")

    lines.append("## Daily summaries")
    for item in daily[:14]:
        date = item.get("date") or item.get("date_time") or item.get("created_at") or item.get("id")
        body = first_text(item, ["summary", "text", "content", "short_summary"])
        lines.append(f"- {date}: {body[:1000]}")
    lines.append("")

    lines.append("## Conversation signals")
    for item in conversations:
        cid = item.get("id")
        short = first_text(item, ["short_summary", "summary"])
        state = item.get("state", "")
        count = item.get("utterances_count", "")
        lines.append(f"- {cid} ({state}, utterances={count}): {short[:900]}")
    lines.append("")

    lines.append("## Facts created/updated in window")
    for item in facts[:250]:
        fid = item.get("id")
        body = first_text(item, ["text", "content", "summary", "title"])
        lines.append(f"- [{fid}] {body[:700]}")
    lines.append("")

    lines.append("## Todos created/alarmed in window")
    for item in todos[:250]:
        tid = item.get("id")
        body = first_text(item, ["text", "content", "title", "summary"])
        status = item.get("status") or item.get("state") or ("completed" if item.get("completed") else "")
        alarm = item.get("alarm_at") or item.get("due_at") or ""
        lines.append(f"- [{tid}] {status} {alarm}: {body[:700]}")
    lines.append("")

    lines.append("## Journals")
    for item in journals[:80]:
        jid = item.get("id")
        body = first_text(item, ["text", "content", "summary", "title"])
        lines.append(f"- [{jid}] {body[:1200]}")
    lines.append("")

    # Keep transcript evidence local without copying transcript text into report artifacts.
    transcript_dir = source_dir / "transcripts"
    lines.append("## Transcript files available")
    md_files = sorted(transcript_dir.glob("*.md"))
    lines.append(f"{len(md_files)} transcript markdown files are available under `{transcript_dir}` for evidence checks.")
    for path in md_files[:80]:
        lines.append(f"- {path.name}")

    out = report_root / "source-digest.md"
    out.write_text("\n".join(lines) + "\n")

    template = report_root / "actionable-report.md"
    if not template.exists():
        template.write_text(default_report_template(window, counts))

    print(out)
    print(template)


def default_report_template(window: dict, counts: dict) -> str:
    return f"""# Weekly Bee Action Report — {window.get('report_date', '')}

Window: {window.get('since_local', '?')} → {window.get('until_local', '?')} {window.get('timezone', 'America/Chicago')}
Source counts: {counts}

## Executive Summary

## Top Pain Points

## Top 3 Actions This Week

## Family / Home / Relationship Action Plan

## Work / Church / Project Action Plan

## Action Packets, Scripts, and Checklists

## Chris Must Decide or Do

## Assistant Can Do Next Without External Side Effects

## Approval Needed

## Blockers / Caveats

"""


def add_markdown_to_doc(doc: Document, markdown: str) -> None:
    in_code = False
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if not line:
            doc.add_paragraph("")
            continue
        if in_code:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Courier New"
            continue
        if line.startswith("# "):
            doc.add_heading(strip_md(line[2:]), 1)
        elif line.startswith("## "):
            doc.add_heading(strip_md(line[3:]), 2)
        elif line.startswith("### "):
            doc.add_heading(strip_md(line[4:]), 3)
        elif line.startswith("#### "):
            doc.add_heading(strip_md(line[5:]), 4)
        elif re.match(r"^\s*[-*] ", line):
            doc.add_paragraph(strip_md(re.sub(r"^\s*[-*] ", "", line)), style="List Bullet")
        elif re.match(r"^\s*\d+\. ", line):
            doc.add_paragraph(strip_md(re.sub(r"^\s*\d+\. ", "", line)), style="List Number")
        elif line.startswith("> "):
            doc.add_paragraph(strip_md(line[2:]))
        else:
            doc.add_paragraph(strip_md(line))


def render(args: argparse.Namespace) -> None:
    report_root = Path(args.report_root).resolve()
    report_date = report_root.name
    md_path = report_root / "actionable-report.md"
    if not md_path.exists():
        raise SystemExit(f"Missing {md_path}")
    out_dir = ROOT / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = Path(args.docx) if args.docx else out_dir / f"bee-weekly-action-report-{report_date}.docx"
    pdf_path = Path(args.pdf) if args.pdf else out_dir / f"bee-weekly-action-report-{report_date}.pdf"

    validate_action_report(md_path)

    if shutil.which("pandoc"):
        subprocess.run(["pandoc", str(md_path), "-o", str(docx_path)], check=True)
    else:
        try:
            from docx import Document
            from docx.shared import Pt
        except Exception as exc:
            marker = report_root / "render-skipped.txt"
            marker.write_text(
                "DOCX/PDF rendering skipped: install pandoc or python-docx. "
                f"Import error: {exc}\n"
            )
            print(marker)
            return

        doc = Document()
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)
        add_markdown_to_doc(doc, md_path.read_text(errors="replace"))
        doc.save(docx_path)

    # Convert through LibreOffice because Telegram accepts PDFs reliably.
    if shutil.which("libreoffice"):
        tmp_out = pdf_path.parent
        tmp_out.mkdir(parents=True, exist_ok=True)
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(tmp_out), str(docx_path)
        ], check=True)
        generated = tmp_out / (docx_path.stem + ".pdf")
        if generated != pdf_path and generated.exists():
            generated.replace(pdf_path)
    elif shutil.which("pandoc"):
        subprocess.run(["pandoc", str(md_path), "-o", str(pdf_path)], check=True)
    else:
        (report_root / "render-warnings.log").write_text("PDF rendering skipped: libreoffice not found.\n")

    print(docx_path)
    if pdf_path.exists():
        print(pdf_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("digest", help="Build source-digest.md and report template")
    p.add_argument("--report-root", required=True)
    p.add_argument("--source-dir", required=True)
    p.set_defaults(func=digest)

    p = sub.add_parser("build", help="Build deterministic actionable-report.md from exported Bee data")
    p.add_argument("--report-root", required=True)
    p.add_argument("--source-dir", required=True)
    p.set_defaults(func=build)

    p = sub.add_parser("render", help="Render actionable-report.md to DOCX/PDF")
    p.add_argument("--report-root", required=True)
    p.add_argument("--docx")
    p.add_argument("--pdf")
    p.set_defaults(func=render)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
